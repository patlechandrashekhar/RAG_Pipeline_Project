"""
docx_builder.py
Professional Word document (.docx) generator for ChipAgent.

Entry point:
    build_docx(spec: DocumentSpec, output_path: str) -> str

All formatting is applied via named styles; no per-run overrides except
for inline emphasis (bold/italic/mono) explicitly requested by a block.
"""

from __future__ import annotations

import io
import os
import warnings
from dataclasses import dataclass, field
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class TableSpec:
    headers: list[str]
    rows: list[list[str]]
    column_widths: Optional[list[float]] = None   # in cm
    zebra: bool = True
    header_fill: str = "2E75B6"                   # hex, no #
    caption: str = ""


@dataclass
class ParagraphBlock:
    text: str
    bold: bool = False
    italic: bool = False
    align: str = "left"   # left | center | right | justify


@dataclass
class BulletBlock:
    items: list[str]


@dataclass
class NumberedBlock:
    items: list[str]


@dataclass
class CodeBlock:
    code: str
    language: str = ""
    caption: str = ""


@dataclass
class CalloutBlock:
    text: str
    kind: str = "note"   # note | warning | tip | important
    title: str = ""


@dataclass
class ImageBlock:
    path: str
    width_cm: float = 14.0
    caption: str = ""


Block = (
    ParagraphBlock
    | BulletBlock
    | NumberedBlock
    | TableSpec
    | CodeBlock
    | CalloutBlock
    | ImageBlock
)


@dataclass
class SectionSpec:
    heading: str
    level: int = 1             # 1, 2, or 3
    blocks: list = field(default_factory=list)


@dataclass
class DocumentSpec:
    title: str
    subtitle: str = ""
    author: str = ""
    date: str = ""
    confidentiality_label: str = ""
    page_size: str = "letter"  # "letter" or "a4"
    include_toc: bool = True
    include_title_page: bool = True
    header_text: str = ""      # shown in page header (defaults to title)
    sections: list[SectionSpec] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

LETTER_W = Cm(21.59)
LETTER_H = Cm(27.94)
A4_W = Cm(21.0)
A4_H = Cm(29.7)

MARGIN = Cm(2.5)

FONT_BODY = "Calibri"
FONT_HEAD = "Calibri"
FONT_CODE = "Courier New"

COLOR_H1 = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
COLOR_H2 = RGBColor(0x2E, 0x75, 0xB6)   # medium blue
COLOR_H3 = RGBColor(0x2E, 0x75, 0xB6)
COLOR_BODY = RGBColor(0x00, 0x00, 0x00)

CALLOUT_COLORS = {
    "note":      ("D6E4F0", "1A5276", "NOTE"),
    "warning":   ("FDEBD0", "784212", "WARNING"),
    "tip":       ("D5F5E3", "1E5631", "TIP"),
    "important": ("F9EBEA", "78281F", "IMPORTANT"),
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    """Apply background fill to a table cell via OOXML shading."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    tc_pr = cell._tc.get_or_add_tcPr()
    # Remove any existing shd element
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    tc_pr.append(shading)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    """Set internal cell padding (in twentieths of a point)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    # Remove existing margins
    for existing in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(existing)
    tc_pr.append(mar)


def _add_field(paragraph, field_code: str) -> None:
    """Insert a Word field (e.g. PAGE, NUMPAGES, TOC) into a paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _para_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.15) -> None:
    """Set paragraph spacing. before/after in pt; line is multiplier."""
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = Pt(12 * line)  # approx 11pt * 1.15


def _align_map(align: str) -> WD_ALIGN_PARAGRAPH:
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(align.lower(), WD_ALIGN_PARAGRAPH.LEFT)


def _safe_text(value) -> str:
    """Ensure value is a non-None string."""
    return "" if value is None else str(value)


def _add_inline_markdown(paragraph, text: str) -> None:
    """
    Add text to a paragraph, parsing inline markdown into styled runs.
    Handles **bold**, *italic*, `monospace`, and ***bold+italic***.
    """
    import re

    # Token pattern: ***bold+italic***, **bold**, *italic*, `code`
    pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    pos = 0
    for m in pattern.finditer(text):
        # Plain text before this match
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])

        triple, double, single, code = m.group(2), m.group(3), m.group(4), m.group(5)
        if triple:
            run = paragraph.add_run(triple)
            run.bold = True
            run.italic = True
        elif double:
            run = paragraph.add_run(double)
            run.bold = True
        elif single:
            run = paragraph.add_run(single)
            run.italic = True
        elif code:
            run = paragraph.add_run(code)
            run.font.name = FONT_CODE
            run.font.size = Pt(9)

        pos = m.end()

    # Remaining plain text
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------------------------------------------------------------------------
# Style definitions
# ---------------------------------------------------------------------------

def _ensure_style(document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH, base: str = "Normal"):
    """Get or create a named paragraph style; return the style object."""
    if name in document.styles:
        return document.styles[name]
    style = document.styles.add_style(name, style_type)
    try:
        style.base_style = document.styles[base]
    except KeyError:
        pass
    return style


def define_styles(document: Document) -> None:
    """Create or configure all named styles used by this builder."""

    # ── Normal ──────────────────────────────────────────────────────────────
    normal = document.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_BODY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = Pt(12 * 1.15)

    # ── Heading 1 ───────────────────────────────────────────────────────────
    h1 = document.styles["Heading 1"]
    h1.font.name = FONT_HEAD
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_H1
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # ── Heading 2 ───────────────────────────────────────────────────────────
    h2 = document.styles["Heading 2"]
    h2.font.name = FONT_HEAD
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_H2
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # ── Heading 3 ───────────────────────────────────────────────────────────
    h3 = document.styles["Heading 3"]
    h3.font.name = FONT_HEAD
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = COLOR_H3
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True

    # ── Caption ─────────────────────────────────────────────────────────────
    cap = _ensure_style(document, "ChipCaption", base="Normal")
    cap.font.name = FONT_BODY
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Code ────────────────────────────────────────────────────────────────
    code_style = _ensure_style(document, "ChipCode", base="Normal")
    code_style.font.name = FONT_CODE
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.left_indent = Cm(0.3)

    # ── Callout ─────────────────────────────────────────────────────────────
    callout_style = _ensure_style(document, "ChipCallout", base="Normal")
    callout_style.font.name = FONT_BODY
    callout_style.font.size = Pt(10)
    callout_style.paragraph_format.left_indent = Cm(0.5)
    callout_style.paragraph_format.right_indent = Cm(0.5)
    callout_style.paragraph_format.space_before = Pt(2)
    callout_style.paragraph_format.space_after = Pt(2)

    # ── List bullet ─────────────────────────────────────────────────────────
    try:
        lb = document.styles["List Bullet"]
        lb.font.name = FONT_BODY
        lb.font.size = Pt(11)
    except KeyError:
        pass

    # ── List number ─────────────────────────────────────────────────────────
    try:
        ln = document.styles["List Number"]
        ln.font.name = FONT_BODY
        ln.font.size = Pt(11)
    except KeyError:
        pass

    # ── Title page title ────────────────────────────────────────────────────
    tp_title = _ensure_style(document, "ChipTitle", base="Normal")
    tp_title.font.name = FONT_HEAD
    tp_title.font.size = Pt(28)
    tp_title.font.bold = True
    tp_title.font.color.rgb = COLOR_H1
    tp_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp_title.paragraph_format.space_before = Pt(80)
    tp_title.paragraph_format.space_after = Pt(8)

    # ── Title page subtitle ─────────────────────────────────────────────────
    tp_sub = _ensure_style(document, "ChipSubtitle", base="Normal")
    tp_sub.font.name = FONT_HEAD
    tp_sub.font.size = Pt(16)
    tp_sub.font.color.rgb = COLOR_H2
    tp_sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp_sub.paragraph_format.space_before = Pt(4)
    tp_sub.paragraph_format.space_after = Pt(40)

    # ── Title page meta ─────────────────────────────────────────────────────
    tp_meta = _ensure_style(document, "ChipMeta", base="Normal")
    tp_meta.font.name = FONT_BODY
    tp_meta.font.size = Pt(11)
    tp_meta.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    tp_meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp_meta.paragraph_format.space_before = Pt(2)
    tp_meta.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def set_page_layout(document: Document, page_size: str = "letter") -> None:
    """Set page size and margins for all sections."""
    if page_size.lower() == "a4":
        w, h = A4_W, A4_H
    else:
        w, h = LETTER_W, LETTER_H

    for section in document.sections:
        section.page_width = w
        section.page_height = h
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------

def add_title_page(document: Document, spec: DocumentSpec) -> None:
    """Insert a title page followed by a page break."""
    document.add_paragraph(_safe_text(spec.title), style="ChipTitle")

    if spec.subtitle:
        document.add_paragraph(_safe_text(spec.subtitle), style="ChipSubtitle")

    if spec.author:
        p = document.add_paragraph(style="ChipMeta")
        p.add_run("Author: ").bold = True
        p.add_run(_safe_text(spec.author))

    if spec.date:
        p = document.add_paragraph(style="ChipMeta")
        p.add_run("Date: ").bold = True
        p.add_run(_safe_text(spec.date))

    if spec.confidentiality_label:
        p = document.add_paragraph(style="ChipMeta")
        run = p.add_run(_safe_text(spec.confidentiality_label).upper())
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # Horizontal rule via bottom border on last meta paragraph
    last_p = document.paragraphs[-1]
    pPr = last_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Page break
    document.add_page_break()


# ---------------------------------------------------------------------------
# Table of Contents
# ---------------------------------------------------------------------------

def add_toc(document: Document) -> None:
    """Insert a TOC heading and the Word TOC field (requires Word to update)."""
    document.add_heading("Table of Contents", level=1)
    p = document.add_paragraph()
    _add_field(p, ' TOC \\o "1-3" \\h \\z \\u ')
    p.paragraph_format.space_after = Pt(6)
    document.add_page_break()


# ---------------------------------------------------------------------------
# Header / footer
# ---------------------------------------------------------------------------

def add_header_footer(document: Document, spec: DocumentSpec) -> None:
    """Add document header (title) and footer (Page X of Y)."""
    header_text = _safe_text(spec.header_text or spec.title)

    for section in document.sections:
        section.different_first_page_header_footer = True

        # ── Header ──────────────────────────────────────────────────────────
        hdr = section.header
        hdr.is_linked_to_previous = False
        hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(header_text)
        run.font.name = FONT_BODY
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Bottom border on header paragraph
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

        # ── Footer ──────────────────────────────────────────────────────────
        ftr = section.footer
        ftr.is_linked_to_previous = False
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_pre = fp.add_run("Page ")
        run_pre.font.name = FONT_BODY
        run_pre.font.size = Pt(9)
        run_pre.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        _add_field(fp, " PAGE ")

        run_mid = fp.add_run(" of ")
        run_mid.font.name = FONT_BODY
        run_mid.font.size = Pt(9)
        run_mid.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        _add_field(fp, " NUMPAGES ")

        if spec.confidentiality_label:
            fp.add_run("  |  ").font.size = Pt(9)
            conf_run = fp.add_run(_safe_text(spec.confidentiality_label))
            conf_run.font.size = Pt(9)
            conf_run.font.bold = True
            conf_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


# ---------------------------------------------------------------------------
# Content block renderers
# ---------------------------------------------------------------------------

def _add_paragraph_block(document: Document, block: ParagraphBlock) -> None:
    p = document.add_paragraph(style="Normal")
    p.alignment = _align_map(block.align)
    if block.bold or block.italic:
        # Whole-paragraph formatting takes precedence
        run = p.add_run(_safe_text(block.text))
        run.bold = block.bold
        run.italic = block.italic
    else:
        _add_inline_markdown(p, _safe_text(block.text))
    _para_spacing(p)


def _add_bullets(document: Document, block: BulletBlock) -> None:
    for item in block.items:
        p = document.add_paragraph(style="List Bullet")
        _add_inline_markdown(p, _safe_text(item))
        _para_spacing(p, before=0, after=3)


def _add_numbered(document: Document, block: NumberedBlock) -> None:
    for item in block.items:
        p = document.add_paragraph(style="List Number")
        _add_inline_markdown(p, _safe_text(item))
        _para_spacing(p, before=0, after=3)


def _add_table(document: Document, spec: TableSpec) -> None:
    """Render a professional table with shaded header row."""
    n_cols = len(spec.headers)
    if n_cols == 0:
        return

    table = document.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths
    if spec.column_widths and len(spec.column_widths) == n_cols:
        for i, col in enumerate(table.columns):
            col.width = Cm(spec.column_widths[i])
    else:
        # Distribute evenly across usable width (21.59 - 5 cm margins = ~16.59)
        avail = Cm(16.59)
        col_w = avail // n_cols
        for col in table.columns:
            col.width = col_w

    # Header row
    hdr_row = table.rows[0]
    hdr_row.height = Pt(20)
    for i, cell in enumerate(hdr_row.cells):
        _set_cell_bg(cell, spec.header_fill)
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(_safe_text(spec.headers[i]))
        run.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    ZEBRA_FILL = "EBF5FB"
    for row_idx, row_data in enumerate(spec.rows):
        row = table.add_row()
        for col_idx, cell in enumerate(row.cells):
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if spec.zebra and row_idx % 2 == 1:
                _set_cell_bg(cell, ZEBRA_FILL)
            p = cell.paragraphs[0]
            p.clear()
            val = _safe_text(row_data[col_idx]) if col_idx < len(row_data) else ""
            run = p.add_run(val)
            run.font.name = FONT_BODY
            run.font.size = Pt(10)

    if spec.caption:
        document.add_paragraph(_safe_text(spec.caption), style="ChipCaption")

    # Spacing after table
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_codeblock(document: Document, block: CodeBlock) -> None:
    """Render a shaded, monospace code block."""
    if block.caption:
        cap = document.add_paragraph(style="ChipCaption")
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.runs[0].text if cap.runs else None
        cap.clear()
        run = cap.add_run(
            (_safe_text(block.language).upper() + " — " if block.language else "")
            + _safe_text(block.caption)
        )
        run.font.italic = True
        run.font.size = Pt(9)

    lines = _safe_text(block.code).splitlines() or [""]
    for i, line in enumerate(lines):
        p = document.add_paragraph(style="ChipCode")
        run = p.add_run(line)
        # Apply light grey background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F3F4")
        for existing in pPr.findall(qn("w:shd")):
            pPr.remove(existing)
        pPr.append(shd)
        # Add borders on first and last line for visual block
        if i == 0 or i == len(lines) - 1:
            pBdr = OxmlElement("w:pBdr")
            if i == 0:
                top = OxmlElement("w:top")
                top.set(qn("w:val"), "single")
                top.set(qn("w:sz"), "4")
                top.set(qn("w:color"), "CCCCCC")
                pBdr.append(top)
            if i == len(lines) - 1:
                bot = OxmlElement("w:bottom")
                bot.set(qn("w:val"), "single")
                bot.set(qn("w:sz"), "4")
                bot.set(qn("w:color"), "CCCCCC")
                pBdr.append(bot)
            pPr.append(pBdr)

    # Spacer after
    document.add_paragraph().paragraph_format.space_after = Pt(6)


def _add_callout(document: Document, block: CalloutBlock) -> None:
    """Render a colored callout/admonition box."""
    fill_hex, text_hex, default_title = CALLOUT_COLORS.get(
        block.kind.lower(), CALLOUT_COLORS["note"]
    )
    label = _safe_text(block.title) if block.title else default_title

    # Title line
    tp = document.add_paragraph(style="ChipCallout")
    tp.clear()
    title_run = tp.add_run(f"  {label}  ")
    title_run.bold = True
    title_run.font.name = FONT_BODY
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor(
        int(text_hex[0:2], 16), int(text_hex[2:4], 16), int(text_hex[4:6], 16)
    )
    _apply_paragraph_fill(tp, fill_hex)

    # Body line(s)
    for line in _safe_text(block.text).splitlines() or [""]:
        bp = document.add_paragraph(style="ChipCallout")
        bp.clear()
        body_run = bp.add_run(f"  {line}")
        body_run.font.name = FONT_BODY
        body_run.font.size = Pt(10)
        _apply_paragraph_fill(bp, fill_hex)

    # Spacer
    document.add_paragraph().paragraph_format.space_after = Pt(6)


def _apply_paragraph_fill(paragraph, hex_color: str) -> None:
    """Apply background shading to an entire paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    for existing in pPr.findall(qn("w:shd")):
        pPr.remove(existing)
    pPr.append(shd)


def _add_image(document: Document, block: ImageBlock) -> None:
    """Add an inline image; skip gracefully if path not found."""
    path = _safe_text(block.path)
    if not os.path.isfile(path):
        warn_p = document.add_paragraph(style="ChipCallout")
        warn_p.add_run(f"  [IMAGE NOT FOUND: {path}]").italic = True
        _apply_paragraph_fill(warn_p, CALLOUT_COLORS["warning"][0])
        return

    try:
        document.add_picture(path, width=Cm(block.width_cm))
        # Center the picture paragraph
        last_p = document.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as exc:
        warn_p = document.add_paragraph(style="Normal")
        warn_p.add_run(f"[Could not embed image '{path}': {exc}]").italic = True

    if block.caption:
        document.add_paragraph(_safe_text(block.caption), style="ChipCaption")


# ---------------------------------------------------------------------------
# Section dispatcher
# ---------------------------------------------------------------------------

_HEADING_LEVEL = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}


def _render_section(document: Document, sec: SectionSpec) -> None:
    level = max(1, min(3, sec.level))
    document.add_heading(_safe_text(sec.heading), level=level)

    for block in sec.blocks:
        if isinstance(block, ParagraphBlock):
            _add_paragraph_block(document, block)
        elif isinstance(block, BulletBlock):
            _add_bullets(document, block)
        elif isinstance(block, NumberedBlock):
            _add_numbered(document, block)
        elif isinstance(block, TableSpec):
            _add_table(document, block)
        elif isinstance(block, CodeBlock):
            _add_codeblock(document, block)
        elif isinstance(block, CalloutBlock):
            _add_callout(document, block)
        elif isinstance(block, ImageBlock):
            _add_image(document, block)
        else:
            # Fallback: render as plain paragraph
            document.add_paragraph(_safe_text(str(block)), style="Normal")


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def build_docx(spec: DocumentSpec, output_path: str) -> str:
    """
    Build a professional Word document from a DocumentSpec.

    Args:
        spec: Fully populated DocumentSpec describing the document content.
        output_path: File path where the .docx will be written.

    Returns:
        Absolute path to the created file.
    """
    document = Document()

    # 1. Page layout
    set_page_layout(document, spec.page_size)

    # 2. Styles
    define_styles(document)

    # 3. Title page
    if spec.include_title_page:
        add_title_page(document, spec)

    # 4. Header / footer (applies to all sections including post-title)
    add_header_footer(document, spec)

    # 5. Table of contents
    if spec.include_toc:
        add_toc(document)

    # 6. Content sections
    for sec in spec.sections:
        _render_section(document, sec)

    # 7. Save
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    document.save(output_path)
    return output_path


def build_docx_bytes(spec: DocumentSpec) -> bytes:
    """
    Build a document and return the raw bytes (useful for Streamlit downloads).

    Args:
        spec: Fully populated DocumentSpec.

    Returns:
        Raw .docx file bytes.
    """
    document = Document()
    set_page_layout(document, spec.page_size)
    define_styles(document)
    if spec.include_title_page:
        add_title_page(document, spec)
    add_header_footer(document, spec)
    if spec.include_toc:
        add_toc(document)
    for sec in spec.sections:
        _render_section(document, sec)

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.read()
